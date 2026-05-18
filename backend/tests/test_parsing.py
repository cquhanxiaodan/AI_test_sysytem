from fastapi.testclient import TestClient
from io import BytesIO

from docx import Document

from app.main import app
from app.modules.documents.repository import DOCUMENTS
from app.modules.parsing.service import CHUNKS, TASKS


client = TestClient(app)


def setup_function() -> None:
    from uuid import uuid4

    from app.modules.admin import service as admin_service

    admin_service.CONFIG = admin_service.DEFAULT_CONFIG.model_copy(deep=True)
    admin_service.get_settings().repository_backend = "memory"
    admin_service.get_settings().system_config_path = f"/tmp/monkeycode-test-system-config-{uuid4()}.json"
    DOCUMENTS.clear()
    TASKS.clear()
    CHUNKS.clear()


def auth_headers() -> dict[str, str]:
    response = client.post("/api/auth/login", json={"username": "admin", "password": "admin123"})
    assert response.status_code == 200
    return {"Authorization": f"Bearer {response.json()['access_token']}"}


def upload_demo(headers: dict[str, str]) -> str:
    response = client.post(
        "/api/documents/upload",
        headers=headers,
        data={"project_id": "project-g99-rfid"},
        files={"file": ("DNBSEQ-G99 RFID验证方案.txt", b"3.1 RFID read test\nstep one", "text/plain")},
    )
    assert response.status_code == 200
    return response.json()["document"]["id"]


def test_parse_document_creates_chunks() -> None:
    headers = auth_headers()
    document_id = upload_demo(headers)

    response = client.post(f"/api/parsing/documents/{document_id}/parse", headers=headers)

    assert response.status_code == 200
    assert response.json()["status"] == "succeeded"
    assert len(response.json()["chunks"]) == 2


def test_label_extraction_updates_high_confidence_labels() -> None:
    headers = auth_headers()
    document_id = upload_demo(headers)

    response = client.post(f"/api/parsing/documents/{document_id}/extract-labels", headers=headers)

    assert response.status_code == 200

    document = client.get(f"/api/documents/{document_id}", headers=headers).json()
    assert document["labels"]["product_model"] == "DNBSEQ-G99"
    assert document["labels"]["subsystem"] == "电子子系统"
    assert document["labels"]["module"] == "RFID"


def test_label_extraction_merges_ai_labels(monkeypatch) -> None:
    headers = auth_headers()
    document_id = upload_demo(headers)

    def fake_run_json_task(*args, **kwargs):
        return {"labels": {"change_type": "供应商变更"}, "confidence": 0.91, "evidence": "文件内容提到 RFID 验证"}

    monkeypatch.setattr("app.modules.parsing.service.run_json_task", fake_run_json_task)

    response = client.post(f"/api/parsing/documents/{document_id}/extract-labels", headers=headers)

    assert response.status_code == 200
    document = client.get(f"/api/documents/{document_id}", headers=headers).json()
    assert document["labels"]["change_type"] == "供应商变更"


def test_label_extraction_normalizes_ai_rfid_subsystem(monkeypatch) -> None:
    headers = auth_headers()
    document_id = upload_demo(headers)

    def fake_run_json_task(*args, **kwargs):
        return {"labels": {"subsystem": "RFID", "change_type": "供应商变更"}, "confidence": 0.91, "evidence": "文件内容提到 RFID 验证"}

    monkeypatch.setattr("app.modules.parsing.service.run_json_task", fake_run_json_task)

    response = client.post(f"/api/parsing/documents/{document_id}/extract-labels", headers=headers)

    assert response.status_code == 200
    document = client.get(f"/api/documents/{document_id}", headers=headers).json()
    assert document["labels"]["subsystem"] == "电子子系统"
    assert document["labels"]["module"] == "RFID"


def test_parse_docx_extracts_paragraphs_and_tables() -> None:
    headers = auth_headers()
    document = Document()
    document.add_paragraph("3.1 RFID 在机读取测试")
    document.add_paragraph("测试目的：验证 RFID 读取稳定性。")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "步骤"
    table.cell(0, 1).text = "预期结果"
    table.cell(1, 0).text = "执行读取"
    table.cell(1, 1).text = "读取成功"
    buffer = BytesIO()
    document.save(buffer)

    upload = client.post(
        "/api/documents/upload",
        headers=headers,
        data={"project_id": "project-g99-rfid"},
        files={"file": ("RFID验证方案.docx", buffer.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    document_id = upload.json()["document"]["id"]

    response = client.post(f"/api/parsing/documents/{document_id}/parse", headers=headers)

    assert response.status_code == 200
    chunk_text = "\n".join(chunk["text"] for chunk in response.json()["chunks"])
    assert "RFID 在机读取测试" in chunk_text
    assert "执行读取 | 读取成功" in chunk_text
    assert "真实解析器接入后" not in chunk_text
