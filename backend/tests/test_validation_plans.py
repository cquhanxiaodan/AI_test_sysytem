from pathlib import Path

from fastapi.testclient import TestClient
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.table import Table
from docx.text.paragraph import Paragraph

from app.main import app
from app.modules.documents.repository import DOCUMENTS
from app.modules.parsing.service import CHUNKS, TASKS
from app.modules.requirements.service import ANALYSES
from app.modules.risks.service import RISKS
from app.modules.test_items.service import TEST_ITEMS
from app.modules.test_packages.service import TEST_PACKAGES
from app.modules.validation_plans.service import EXPORTS, PLANS


client = TestClient(app)


def setup_function() -> None:
    from uuid import uuid4

    from app.modules.admin import service as admin_service
    from app.modules.ai.service import RUNTIME_AI_CONFIG

    admin_service.CONFIG = admin_service.DEFAULT_CONFIG.model_copy(deep=True)
    admin_service.get_settings().repository_backend = "memory"
    admin_service.get_settings().system_config_path = f"/tmp/monkeycode-test-system-config-{uuid4()}.json"
    RUNTIME_AI_CONFIG.clear()
    DOCUMENTS.clear()
    TASKS.clear()
    CHUNKS.clear()
    TEST_ITEMS.clear()
    TEST_PACKAGES.clear()
    RISKS.clear()
    ANALYSES.clear()
    PLANS.clear()
    EXPORTS.clear()


def auth_headers() -> dict[str, str]:
    response = client.post("/api/auth/login", json={"username": "admin", "password": "admin123"})
    assert response.status_code == 200
    return {"Authorization": f"Bearer {response.json()['access_token']}"}


def seed_assets(headers: dict[str, str]) -> None:
    upload = client.post(
        "/api/documents/upload",
        headers=headers,
        data={"project_id": "project-g99-rfid"},
        files={"file": ("RFID验证方案.txt", b"RFID", "text/plain")},
    )
    document_id = upload.json()["document"]["id"]
    client.post(f"/api/test-items/split/{document_id}", headers=headers)
    package = client.post("/api/test-packages/generate-rfid-supplier-change?project_id=project-g99-rfid", headers=headers).json()
    client.post(f"/api/test-packages/{package['id']}/publish", headers=headers)
    risks = client.post(
        "/api/risks/parse",
        headers=headers,
        json={"project_id": "project-g99-rfid", "source_type": "jira", "content": "title\nRFID读取失败\n"},
    ).json()["items"]
    client.post("/api/risks/bulk-publish", headers=headers, json={"risk_ids": [item["id"] for item in risks]})


def create_analysis(headers: dict[str, str]) -> str:
    response = client.post(
        "/api/requirement-analyses/local",
        headers=headers,
        json={"project_id": "project-g99-rfid", "description": "DNBSEQ-G99 引入二供供应商康奈特 RFID"},
    )
    assert response.status_code == 200
    return response.json()["id"]


def update_recommendation_status(headers: dict[str, str], analysis_id: str, index: int, review_status: str) -> dict:
    analysis = client.get(f"/api/requirement-analyses/{analysis_id}", headers=headers).json()
    recommendation_id = analysis["recommendations"][index]["id"]
    response = client.patch(
        f"/api/requirement-analyses/{analysis_id}/recommendations/{recommendation_id}",
        headers=headers,
        json={"review_status": review_status},
    )
    assert response.status_code == 200
    return response.json()["recommendations"][index]


def test_create_plan_from_project_batches_all_analyses(monkeypatch) -> None:
    monkeypatch.setattr("app.modules.validation_plans.service.run_json_task", lambda *args, **kwargs: None)
    headers = auth_headers()
    seed_assets(headers)
    first_analysis_id = create_analysis(headers)
    second_analysis_id = create_analysis(headers)
    update_recommendation_status(headers, first_analysis_id, 0, "confirmed")
    update_recommendation_status(headers, second_analysis_id, 0, "confirmed")

    created = client.post("/api/validation-plans", headers=headers, json={"project_id": "project-g99-rfid"})

    assert created.status_code == 200
    plan = created.json()
    assert plan["template_version"] == "validation-plan-v1"
    assert len(plan["requirement_analysis_ids"]) == 2
    assert len(plan["items"]) > 0

    checked = client.post(f"/api/validation-plans/{plan['id']}/check", headers=headers)
    assert checked.status_code == 200
    assert "blocking" in checked.json()

    exported = client.post(f"/api/validation-plans/{plan['id']}/export", headers=headers)
    assert exported.status_code == 200
    export_record = exported.json()
    assert export_record["filename"].endswith(".docx")
    assert Path(export_record["storage_path"]).exists()
    document = Document(export_record["storage_path"])
    expected_title = export_record["filename"].removesuffix(".docx")
    header_cells = [cell.text for section in document.sections for table in section.header.tables for cell in table.rows[0].cells]
    assert document.core_properties.title == expected_title
    assert expected_title in header_cells
    assert all("DNBSEQ-G99 ECR4.1康奈特RFID验证方案" not in cell for cell in header_cells)

    downloaded = client.get(export_record["download_url"], headers=headers)
    assert downloaded.status_code == 200
    assert downloaded.content.startswith(b"PK")


def test_create_plan_from_single_analysis_still_works() -> None:
    headers = auth_headers()
    seed_assets(headers)
    analysis_id = create_analysis(headers)
    update_recommendation_status(headers, analysis_id, 0, "confirmed")

    created = client.post("/api/validation-plans", headers=headers, json={"requirement_analysis_id": analysis_id})

    assert created.status_code == 200
    plan = created.json()
    assert len(plan["requirement_analysis_ids"]) == 1


def test_validation_plan_check_merges_ai_messages(monkeypatch) -> None:
    headers = auth_headers()
    seed_assets(headers)
    analysis_id = create_analysis(headers)
    update_recommendation_status(headers, analysis_id, 0, "confirmed")
    created = client.post("/api/validation-plans", headers=headers, json={"project_id": "project-g99-rfid"})
    plan_id = created.json()["id"]

    def fake_run_json_task(*args, **kwargs):
        return {"blocking": [], "warnings": ["AI 提示需确认 RFID DUT 批次。"], "suggestions": ["AI 建议补充异常恢复记录。"]}

    monkeypatch.setattr("app.modules.validation_plans.service.run_json_task", fake_run_json_task)

    checked = client.post(f"/api/validation-plans/{plan_id}/check", headers=headers)

    assert checked.status_code == 200
    result = checked.json()
    assert "AI 提示需确认 RFID DUT 批次。" in result["warnings"]
    assert "AI 建议补充异常恢复记录。" in result["suggestions"]


def test_validation_plan_uses_only_confirmed_recommendations() -> None:
    headers = auth_headers()
    seed_assets(headers)
    analysis_id = create_analysis(headers)
    confirmed_item = update_recommendation_status(headers, analysis_id, 0, "confirmed")
    excluded_item = update_recommendation_status(headers, analysis_id, 1, "excluded")

    created = client.post("/api/validation-plans", headers=headers, json={"project_id": "project-g99-rfid"})

    assert created.status_code == 200
    titles = {item["title"] for item in created.json()["items"]}
    assert confirmed_item["title"] in titles
    assert excluded_item["title"] not in titles
    assert len(titles) == 1


def test_validation_plan_excludes_all_pending_recommendations() -> None:
    headers = auth_headers()
    seed_assets(headers)
    create_analysis(headers)

    created = client.post("/api/validation-plans", headers=headers, json={"project_id": "project-g99-rfid"})

    assert created.status_code == 200
    assert created.json()["items"] == []


def test_validation_plan_items_carry_requirement_fields() -> None:
    headers = auth_headers()
    seed_assets(headers)
    analysis_id = create_analysis(headers)
    update_recommendation_status(headers, analysis_id, 0, "confirmed")

    created = client.post("/api/validation-plans", headers=headers, json={"project_id": "project-g99-rfid"})

    assert created.status_code == 200
    item = created.json()["items"][0]
    assert item["objective"]
    assert item["method"]
    assert item["record_template"]
    assert item["title"] in item["compliance_bug_info"]


def test_validation_plan_items_reuse_full_test_item_content() -> None:
    headers = auth_headers()
    content = """
测试项目列表
测试项目
整机安装适配测试
整机安装适配测试
测试目的/测试标准
检验RFID结构适配安装至整机是否存在异常。
测试方法/原理
整机安装适配。
测试工具
DNBSEQ-G99
RFID 标签
测试步骤
RFID检查；
RFID安装；
整机外壳安装；
测试连接图或照片
整机安装照片见附件 A。
测试记录
记录结构干涉、线缆干涉和接头状态。
需求符合性和BUG信息
通过时记录需求符合，失败时登记 BUG 编号。
""".strip()
    upload = client.post(
        "/api/documents/upload",
        headers=headers,
        data={"project_id": "project-g99-rfid"},
        files={"file": ("DNBSEQ-G99 RFID验证方案.txt", content.encode("utf-8"), "text/plain")},
    )
    document_id = upload.json()["document"]["id"]
    client.post(f"/api/parsing/documents/{document_id}/parse", headers=headers)
    split = client.post(f"/api/test-items/split/{document_id}", headers=headers)
    item_id = split.json()["items"][0]["id"]
    package = client.post("/api/test-packages/generate-rfid-supplier-change?project_id=project-g99-rfid", headers=headers).json()
    client.post(f"/api/test-packages/{package['id']}/publish", headers=headers)

    analysis = create_analysis(headers)
    update_recommendation_status(headers, analysis, 0, "confirmed")
    created = client.post("/api/validation-plans", headers=headers, json={"project_id": "project-g99-rfid"})

    assert created.status_code == 200
    item = created.json()["items"][0]
    assert item_id
    assert item["tools"] == ["DNBSEQ-G99", "RFID 标签"]
    assert item["steps"] == ["RFID检查", "RFID安装", "整机外壳安装"]
    assert item["connection_media"] == "整机安装照片见附件 A。"
    assert item["compliance_bug_info"] == "通过时记录需求符合，失败时登记 BUG 编号。"
    assert "测试目的/测试标准" in item["source_section_text"]


def test_validation_plan_export_uses_structured_test_item_headings() -> None:
    headers = auth_headers()
    seed_assets(headers)
    analysis_id = create_analysis(headers)
    update_recommendation_status(headers, analysis_id, 0, "confirmed")
    created = client.post("/api/validation-plans", headers=headers, json={"project_id": "project-g99-rfid"})
    plan_id = created.json()["id"]
    expected_title = created.json()["items"][0]["title"]

    exported = client.post(f"/api/validation-plans/{plan_id}/export", headers=headers)

    assert exported.status_code == 200
    document = Document(exported.json()["storage_path"])
    heading_texts = [paragraph.text for paragraph in document.paragraphs if paragraph.style.name.startswith("Heading")]
    assert expected_title in heading_texts
    assert "需求符合性和BUG信息" in heading_texts
    assert all(not text.startswith("3.") for text in heading_texts)
    for paragraph in document.paragraphs:
        if paragraph.text in {expected_title, "测试工具", "测试记录", "需求符合性和BUG信息"}:
            assert paragraph.paragraph_format.left_indent == Pt(0)
            assert paragraph.paragraph_format.first_line_indent == Pt(0)
            assert paragraph._p.pPr.ind.get(qn("w:leftChars")) == "0"
            assert paragraph._p.pPr.ind.get(qn("w:firstLineChars")) == "0"


def test_validation_plan_export_body_paragraphs_have_consistent_indent() -> None:
    headers = auth_headers()
    seed_assets(headers)
    analysis_id = create_analysis(headers)
    recommendation = update_recommendation_status(headers, analysis_id, 0, "confirmed")
    created = client.post("/api/validation-plans", headers=headers, json={"project_id": "project-g99-rfid"})
    plan_id = created.json()["id"]

    exported = client.post(f"/api/validation-plans/{plan_id}/export", headers=headers)

    assert exported.status_code == 200
    document = Document(exported.json()["storage_path"])
    expected_body_lines = [recommendation["objective"], recommendation["method"]]
    for paragraph in document.paragraphs:
        if paragraph.text in expected_body_lines or paragraph.text.startswith("1. "):
            assert paragraph.paragraph_format.left_indent == Pt(0)
            assert paragraph._p.pPr.ind.get(qn("w:leftChars")) == "0"
            assert paragraph._p.pPr.ind.get(qn("w:firstLine")) == "480"
            assert paragraph._p.pPr.ind.get(qn("w:firstLineChars")) is None


def test_validation_plan_export_renders_structured_tables() -> None:
    headers = auth_headers()
    seed_assets(headers)
    analysis_id = create_analysis(headers)
    update_recommendation_status(headers, analysis_id, 0, "confirmed")
    created = client.post("/api/validation-plans", headers=headers, json={"project_id": "project-g99-rfid"})
    plan_id = created.json()["id"]

    exported = client.post(f"/api/validation-plans/{plan_id}/export", headers=headers)

    assert exported.status_code == 200
    document = Document(exported.json()["storage_path"])
    table_headers = [tuple(cell.text for cell in table.rows[0].cells) for table in document.tables]
    table_rows = [tuple(cell.text for cell in row.cells) for table in document.tables for row in table.rows]
    assert ("序号", "名称", "设备型号", "制造商", "设备编码", "校准有效期") in table_headers
    assert any("记录项目" in "".join(row) and "判定标准" in "".join(row) for row in table_rows)
    record_table = next(
        table
        for table in document.tables
        if len(table.rows) > 3 and any("记录项目" in cell.text for cell in table.rows[3].cells)
    )
    assert len(record_table.columns) == 9
    assert record_table.rows[0].cells[0]._tc is record_table.rows[0].cells[1]._tc
    assert record_table.rows[0].cells[4]._tc is record_table.rows[0].cells[5]._tc
    assert record_table.rows[0].cells[6]._tc is record_table.rows[0].cells[8]._tc
    assert ("序号", "需求编号/DFMEA编号/风险管理编号", "需求描述", "测试结论", "备注") in table_headers
    assert ("序号", "问题描述", "涉及需求编号", "BUG编号（JIRA系统）", "RPN", "Bug解决状态") in table_headers
    body_items = list(iter_document_body_items(document))
    compliance_index = find_table_index(body_items, ("序号", "需求编号/DFMEA编号/风险管理编号", "需求描述", "测试结论", "备注"))
    bug_index = find_table_index(body_items, ("序号", "问题描述", "涉及需求编号", "BUG编号（JIRA系统）", "RPN", "Bug解决状态"))
    assert isinstance(body_items[compliance_index + 1], Paragraph)
    assert body_items[compliance_index + 1].text == ""
    assert compliance_index + 2 == bug_index


def test_validation_plan_export_renders_top_level_tables() -> None:
    headers = auth_headers()
    seed_assets(headers)
    analysis_id = create_analysis(headers)
    update_recommendation_status(headers, analysis_id, 0, "confirmed")
    created = client.post("/api/validation-plans", headers=headers, json={"project_id": "project-g99-rfid"})
    plan_id = created.json()["id"]

    exported = client.post(f"/api/validation-plans/{plan_id}/export", headers=headers)

    assert exported.status_code == 200
    document = Document(exported.json()["storage_path"])
    table_headers = [tuple(cell.text for cell in table.rows[0].cells) for table in document.tables]
    assert ("序号", "名称", "型号", "物料编码/版本", "制造商", "物料编号", "测试数量") in table_headers
    assert ("序号", "名称", "编号", "版本", "创建人", "时间") in table_headers
    assert ("序号", "测试项目", "对应需求编号/DFMEA编号/风险管理编号/测试目的", "样本量", "预估测试用时（h）", "备注") in table_headers
    table = next(table for table in document.tables if tuple(cell.text for cell in table.rows[0].cells) == table_headers[-1])
    for cell in table.rows[0].cells:
        assert cell.vertical_alignment == WD_CELL_VERTICAL_ALIGNMENT.CENTER
        assert cell.paragraphs[0].alignment == WD_PARAGRAPH_ALIGNMENT.CENTER
        assert cell.paragraphs[0].paragraph_format.first_line_indent == Pt(0)
        assert cell.paragraphs[0].paragraph_format.left_indent == Pt(0)
        indent = cell.paragraphs[0]._p.pPr.ind
        assert indent.get(qn("w:firstLineChars")) == "0"
        assert indent.get(qn("w:leftChars")) == "0"


def test_validation_plan_export_preserves_reference_template_styles() -> None:
    headers = auth_headers()
    seed_assets(headers)
    analysis_id = create_analysis(headers)
    update_recommendation_status(headers, analysis_id, 0, "confirmed")
    created = client.post("/api/validation-plans", headers=headers, json={"project_id": "project-g99-rfid"})
    plan_id = created.json()["id"]

    exported = client.post(f"/api/validation-plans/{plan_id}/export", headers=headers)

    assert exported.status_code == 200
    document = Document(exported.json()["storage_path"])
    paragraphs = [(paragraph.style.name, paragraph.text) for paragraph in document.paragraphs if paragraph.text.strip()]
    assert ("Normal", "文档履历") in paragraphs
    assert any(style == "Heading 1" and text == "概述" for style, text in paragraphs)
    assert any(style == "Heading 1" and text == "测试项目" for style, text in paragraphs)


def test_validation_plan_export_marks_toc_fields_for_refresh() -> None:
    headers = auth_headers()
    seed_assets(headers)
    analysis_id = create_analysis(headers)
    update_recommendation_status(headers, analysis_id, 0, "confirmed")
    created = client.post("/api/validation-plans", headers=headers, json={"project_id": "project-g99-rfid"})
    plan_id = created.json()["id"]

    exported = client.post(f"/api/validation-plans/{plan_id}/export", headers=headers)

    assert exported.status_code == 200
    document = Document(exported.json()["storage_path"])
    update_fields = document.settings.element.find(qn("w:updateFields"))
    assert update_fields is not None
    assert update_fields.get(qn("w:val")) == "true"
    field_instructions = [node.text for node in document.element.body.iter(qn("w:instrText"))]
    assert any("TOC" in instruction for instruction in field_instructions if instruction)
    assert any(field.get(qn("w:fldCharType")) == "begin" for field in document.element.body.iter(qn("w:fldChar")))
    assert any(field.get(qn("w:fldCharType")) == "separate" for field in document.element.body.iter(qn("w:fldChar")))
    assert any(field.get(qn("w:fldCharType")) == "end" for field in document.element.body.iter(qn("w:fldChar")))


def test_validation_plan_export_rebuilds_toc_with_current_items() -> None:
    headers = auth_headers()
    seed_assets(headers)
    analysis_id = create_analysis(headers)
    update_recommendation_status(headers, analysis_id, 0, "confirmed")
    created = client.post("/api/validation-plans", headers=headers, json={"project_id": "project-g99-rfid"})
    plan_id = created.json()["id"]
    expected_title = created.json()["items"][0]["title"]

    exported = client.post(f"/api/validation-plans/{plan_id}/export", headers=headers)

    assert exported.status_code == 200
    document = Document(exported.json()["storage_path"])
    body_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    toc_title = next(paragraph for paragraph in document.paragraphs if paragraph.text == "目 录")
    toc_item = next(paragraph for paragraph in document.paragraphs if paragraph.text.startswith(f"3.1 {expected_title}"))
    assert toc_title.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER
    assert toc_title.runs[0].bold is True
    assert toc_item.text == f"3.1 {expected_title}\t5"
    assert toc_item.paragraph_format.tab_stops[0].alignment == WD_TAB_ALIGNMENT.RIGHT
    assert f"3.1 {expected_title}" in body_text
    assert "整机安装适配测试" not in body_text


def test_validation_plan_export_uses_source_blocks_for_document_items() -> None:
    headers = auth_headers()
    content = """
测试项目列表
测试项目
整机安装适配测试
整机安装适配测试
测试目的/测试标准
检验RFID结构适配安装至整机是否存在异常。
测试方法/原理
整机安装适配。
测试工具
DNBSEQ-G99
RFID 标签
测试步骤
RFID检查；
RFID安装；
测试记录
记录结构干涉；记录线缆干涉。
需求符合性和BUG信息
通过时记录需求符合，失败时登记 BUG 编号。
""".strip()
    upload = client.post(
        "/api/documents/upload",
        headers=headers,
        data={"project_id": "project-g99-rfid"},
        files={"file": ("DNBSEQ-G99 RFID验证方案.txt", content.encode("utf-8"), "text/plain")},
    )
    document_id = upload.json()["document"]["id"]
    client.post(f"/api/parsing/documents/{document_id}/parse", headers=headers)
    split = client.post(f"/api/test-items/split/{document_id}", headers=headers)
    item_id = split.json()["items"][0]["id"]
    assert split.json()["items"][0]["source_blocks"]
    client.post(f"/api/test-items/{item_id}/confirm", headers=headers)
    package = client.post("/api/test-packages/generate-rfid-supplier-change?project_id=project-g99-rfid", headers=headers).json()
    client.post(f"/api/test-packages/{package['id']}/publish", headers=headers)
    analysis = create_analysis(headers)
    update_recommendation_status(headers, analysis, 0, "confirmed")
    created = client.post("/api/validation-plans", headers=headers, json={"project_id": "project-g99-rfid"})

    exported = client.post(f"/api/validation-plans/{created.json()['id']}/export", headers=headers)

    assert exported.status_code == 200
    document = Document(exported.json()["storage_path"])
    all_cells = [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    assert "DNBSEQ-G99" in all_cells
    assert "RFID 标签" in all_cells
    assert "检验RFID结构适配安装至整机是否存在异常。" in paragraphs


def iter_document_body_items(document: Document):
    for child in document.element.body.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, document)
        elif child.tag.endswith("}tbl"):
            yield Table(child, document)


def find_table_index(items: list, header: tuple[str, ...]) -> int:
    for index, item in enumerate(items):
        if isinstance(item, Table) and tuple(cell.text for cell in item.rows[0].cells) == header:
            return index
    raise AssertionError(f"table header not found: {header}")


def test_validation_plan_status_can_be_updated_and_export_marks_exported() -> None:
    headers = auth_headers()
    seed_assets(headers)
    analysis_id = create_analysis(headers)
    update_recommendation_status(headers, analysis_id, 0, "confirmed")
    created = client.post("/api/validation-plans", headers=headers, json={"project_id": "project-g99-rfid"})
    plan_id = created.json()["id"]

    updated = client.patch(f"/api/validation-plans/{plan_id}/status", headers=headers, json={"status": "reviewing"})

    assert updated.status_code == 200
    assert updated.json()["status"] == "reviewing"

    exported = client.post(f"/api/validation-plans/{plan_id}/export", headers=headers)
    assert exported.status_code == 200
    detail = client.get(f"/api/validation-plans/{plan_id}", headers=headers)
    assert detail.json()["status"] == "exported"


def test_validation_plan_rejects_unknown_status() -> None:
    headers = auth_headers()
    seed_assets(headers)
    analysis_id = create_analysis(headers)
    update_recommendation_status(headers, analysis_id, 0, "confirmed")
    created = client.post("/api/validation-plans", headers=headers, json={"project_id": "project-g99-rfid"})

    updated = client.patch(f"/api/validation-plans/{created.json()['id']}/status", headers=headers, json={"status": "unknown"})

    assert updated.status_code == 400


def test_validation_plan_can_be_deleted() -> None:
    headers = auth_headers()
    seed_assets(headers)
    analysis_id = create_analysis(headers)
    update_recommendation_status(headers, analysis_id, 0, "confirmed")
    created = client.post("/api/validation-plans", headers=headers, json={"project_id": "project-g99-rfid"})
    plan_id = created.json()["id"]

    deleted = client.delete(f"/api/validation-plans/{plan_id}", headers=headers)

    assert deleted.status_code == 204
    detail = client.get(f"/api/validation-plans/{plan_id}", headers=headers)
    assert detail.status_code == 404


def test_validation_plans_can_be_bulk_deleted() -> None:
    headers = auth_headers()
    seed_assets(headers)
    first_analysis_id = create_analysis(headers)
    second_analysis_id = create_analysis(headers)
    update_recommendation_status(headers, first_analysis_id, 0, "confirmed")
    update_recommendation_status(headers, second_analysis_id, 0, "confirmed")
    first_plan = client.post("/api/validation-plans", headers=headers, json={"requirement_analysis_id": first_analysis_id}).json()
    second_plan = client.post("/api/validation-plans", headers=headers, json={"requirement_analysis_id": second_analysis_id}).json()

    deleted = client.post(
        "/api/validation-plans/bulk-delete",
        headers=headers,
        json={"plan_ids": [first_plan["id"], second_plan["id"], "missing-plan"]},
    )

    assert deleted.status_code == 200
    result = deleted.json()
    assert set(result["deleted_ids"]) == {first_plan["id"], second_plan["id"]}
    assert result["skipped"] == [{"plan_id": "missing-plan", "reason": "方案不存在或无访问权限"}]
    remaining = client.get("/api/validation-plans?project_id=project-g99-rfid", headers=headers)
    assert remaining.json() == []


def test_validation_plan_export_directory_is_chosen_per_export(tmp_path) -> None:
    headers = auth_headers()
    export_directory = tmp_path / "exports"
    seed_assets(headers)
    analysis_id = create_analysis(headers)
    update_recommendation_status(headers, analysis_id, 0, "confirmed")
    created = client.post("/api/validation-plans", headers=headers, json={"project_id": "project-g99-rfid"})

    exported = client.post(
        f"/api/validation-plans/{created.json()['id']}/export",
        headers=headers,
        json={"export_directory": str(export_directory)},
    )

    assert exported.status_code == 200
    assert Path(exported.json()["storage_path"]).is_relative_to(export_directory)
    assert Path(exported.json()["storage_path"]).exists()
