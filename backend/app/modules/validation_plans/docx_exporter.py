from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

from app.modules.validation_plans.schemas import ValidationPlanRead


def ensure_default_template(template_path: Path) -> None:
    if template_path.exists():
        return
    template_path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.add_heading("{{ title }}", level=1)
    document.add_heading("文档履历", level=2)
    history = document.add_table(rows=2, cols=5)
    history.style = "Table Grid"
    for index, value in enumerate(["版本", "日期", "新增/修订内容概述", "编制/修订人", "批准人"]):
        history.rows[0].cells[index].text = value
    for index, value in enumerate(["V0.1", "{{ generated_date }}", "AI 生成验证方案草稿", "AI Assistant", "待审核"]):
        history.rows[1].cells[index].text = value

    document.add_heading("1. 概述", level=2)
    document.add_heading("1.1 验证的背景、目的和范围", level=3)
    document.add_paragraph("{{ overview }}")
    document.add_heading("1.2 DUT描述", level=3)
    document.add_paragraph("{{ dut_description }}")
    document.add_heading("1.3 参考文档", level=3)
    document.add_paragraph("{% for reference in reference_documents %}{{ loop.index }}. {{ reference }}\n{% endfor %}")

    document.add_heading("2. 测试项目列表", level=2)
    document.add_paragraph("{% for item in items %}{{ item.sequence }}. {{ item.title }}（{{ item.group }}）\n{% endfor %}")

    document.add_heading("3. 测试项目", level=2)
    document.add_paragraph("{% for item in items %}3.{{ item.sequence }} {{ item.title }}")
    document.add_paragraph("3.{{ item.sequence }}.1 测试目的/测试标准\n{{ item.objective }}")
    document.add_paragraph("3.{{ item.sequence }}.2 测试方法/原理\n{{ item.method }}")
    document.add_paragraph("3.{{ item.sequence }}.3 测试工具\n{% for tool in item.tools %}{{ loop.index }}. {{ tool }}\n{% endfor %}")
    document.add_paragraph("3.{{ item.sequence }}.4 测试步骤\n{% for step in item.steps %}{{ loop.index }}. {{ step }}\n{% endfor %}")
    document.add_paragraph("3.{{ item.sequence }}.5 测试连接图或照片\n{{ item.connection_media }}")
    document.add_paragraph("3.{{ item.sequence }}.6 测试记录\n{{ item.record_template }}")
    document.add_paragraph("3.{{ item.sequence }}.7 需求符合性和BUG信息\n{{ item.compliance_bug_info }}\n依据：{{ item.evidence }}")
    document.add_paragraph("{% if item.source_section_text %}原始测试项目章节\n{{ item.source_section_text }}{% endif %}")
    document.add_paragraph("{% endfor %}")
    document.save(template_path)


def render_validation_plan_docx(plan: ValidationPlanRead, template_path: Path, output_path: Path) -> None:
    ensure_default_template(template_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    Document(template_path).save(output_path)
    rewrite_overview_section(output_path, plan)
    rewrite_summary_tables(output_path, plan)
    rewrite_test_item_section(output_path, plan)


def rewrite_overview_section(output_path: Path, plan: ValidationPlanRead) -> None:
    document = Document(output_path)
    replace_heading_body_with_paragraphs(document, ["1.1 验证的背景、目的和范围", "验证的背景、目的和范围"], [plan.overview])
    document.save(output_path)


def rewrite_summary_tables(output_path: Path, plan: ValidationPlanRead) -> None:
    document = Document(output_path)
    replace_heading_body_with_table(document, ["1.2 DUT描述", "DUT描述"], build_dut_table_rows(plan))
    replace_heading_body_with_table(document, ["1.3 参考文档", "参考文档"], build_reference_table_rows(plan.reference_documents))
    replace_heading_body_with_table(document, ["2. 测试项目列表", "测试项目列表"], build_test_project_table_rows(plan))
    document.save(output_path)


def rewrite_test_item_section(output_path: Path, plan: ValidationPlanRead) -> None:
    document = Document(output_path)
    start_index = find_heading_index(document, ["3. 测试项目", "测试项目"])
    if start_index is None:
        return
    remove_paragraphs_after(document, start_index)
    for item in plan.items:
        document.add_heading(item.title, level=2)
        if has_source_blocks(item):
            append_source_blocks(document, item.source_blocks)
            continue
        add_test_item_subsection(document, item.sequence, 1, "测试目的/测试标准", item.objective)
        add_test_item_subsection(document, item.sequence, 2, "测试方法/原理", item.method)
        add_test_item_heading(document, item.sequence, 3, "测试工具")
        add_tools_table(document, item.tools)
        add_test_item_subsection(document, item.sequence, 4, "测试步骤", numbered_lines(item.steps, "按模板执行并记录过程数据。"))
        add_test_item_subsection(document, item.sequence, 5, "测试连接图或照片", item.connection_media or "待补充")
        add_test_item_heading(document, item.sequence, 6, "测试记录")
        add_record_table(document, item.title, item.record_template)
        add_test_item_heading(document, item.sequence, 7, "需求符合性和BUG信息")
        add_compliance_table(document, item.title, item.compliance_bug_info)
        document.add_paragraph("")
        add_bug_table(document)
    document.save(output_path)


def replace_heading_body_with_paragraphs(document: Document, heading_texts: list[str], values: list[str]) -> None:
    index = find_heading_index(document, heading_texts)
    if index is None:
        return
    clear_heading_body(document, index)
    anchor = document.paragraphs[index]._p
    for value in reversed([value for value in values if value]):
        paragraph = document.add_paragraph(value)
        anchor.addnext(paragraph._p)


def replace_heading_body_with_table(document: Document, heading_texts: list[str], rows: list[list[str]]) -> None:
    index = find_heading_index(document, heading_texts)
    if index is None:
        return
    clear_heading_body(document, index)
    target_table = document.add_table(rows=len(rows), cols=len(rows[0]))
    document.paragraphs[index]._p.addnext(target_table._tbl)
    target_table.style = "Table Grid"
    for row, row_values in zip(target_table.rows, rows):
        set_table_row(row.cells, row_values)


def find_heading_index(document: Document, heading_texts: list[str]) -> int | None:
    targets = set(heading_texts)
    return next((index for index, paragraph in enumerate(document.paragraphs) if paragraph.text.strip() in targets), None)


def clear_heading_body(document: Document, heading_index: int) -> None:
    heading = document.paragraphs[heading_index]
    current = heading._p.getnext()
    while current is not None:
        if current.tag.endswith("}p"):
            paragraph_text = "".join(node.text or "" for node in current.iter() if node.tag.endswith("}t")).strip()
            paragraph = next((item for item in document.paragraphs if item._p is current), None)
            if paragraph is not None and paragraph.style.name.startswith("Heading") and paragraph_text:
                break
        next_element = current.getnext()
        if not current.tag.endswith("}sectPr"):
            current.getparent().remove(current)
        current = next_element


def build_dut_table_rows(plan: ValidationPlanRead) -> list[list[str]]:
    test_objects = sorted({item.title for item in plan.items})
    return [
        ["序号", "名称", "型号", "物料编码/版本", "制造商", "物料编号", "测试数量"],
        ["1", plan.dut_description or "待确认", "/", "/", "/", "/", str(max(len(test_objects), 1))],
    ]


def build_reference_table_rows(references: list[str]) -> list[list[str]]:
    rows = [["序号", "名称", "编号", "版本", "创建人", "时间"]]
    for index, reference in enumerate(references or ["测试规范"], start=1):
        rows.append([str(index), reference, "/", "/", "/", "/"])
    return rows


def build_test_project_table_rows(plan: ValidationPlanRead) -> list[list[str]]:
    rows = [["序号", "测试项目", "对应需求编号/DFMEA编号/风险管理编号/测试目的", "样本量", "预估测试用时（h）", "备注"]]
    for item in plan.items:
        rows.append([str(item.sequence), item.title, item.objective or item.evidence, "1", "/", item.group])
    return rows


def remove_paragraphs_after(document: Document, start_index: int) -> None:
    current = document.paragraphs[start_index]._p.getnext()
    while current is not None:
        next_element = current.getnext()
        if not current.tag.endswith("}sectPr"):
            current.getparent().remove(current)
        current = next_element


def add_test_item_subsection(document: Document, sequence: int, subsection: int, title: str, body: str) -> None:
    add_test_item_heading(document, sequence, subsection, title)
    document.add_paragraph(body or "待补充")


def add_test_item_heading(document: Document, sequence: int, subsection: int, title: str) -> None:
    document.add_heading(title, level=3)


def add_tools_table(document: Document, tools: list[str]) -> None:
    rows = [tool for tool in tools if tool] or ["/"]
    table = document.add_table(rows=len(rows) + 1, cols=6)
    table.style = "Table Grid"
    set_table_row(table.rows[0].cells, ["序号", "名称", "设备型号", "制造商", "设备编码", "校准有效期"])
    for index, tool in enumerate(rows, start=1):
        set_table_row(table.rows[index].cells, [str(index), tool, "/", "/", "/", "/"])


def add_record_table(document: Document, title: str, record_template: str) -> None:
    records = split_record_lines(record_template) or [title]
    table = document.add_table(rows=len(records) + 4, cols=7)
    table.style = "Table Grid"
    set_table_row(table.rows[0].cells, ["环境温度", "", "", "相对湿度", "", "", ""])
    set_table_row(table.rows[1].cells, ["测试时间", "", "", "测试地点", "", "", ""])
    set_table_row(table.rows[2].cells, ["测试人员", "", "", "测试结论", "", "", ""])
    set_table_row(table.rows[3].cells, ["序号", "记录项目", "记录数据（单位）", "处理结果（单位）", "判定标准", "结果", "备注"])
    for index, record in enumerate(records, start=1):
        set_table_row(table.rows[index + 3].cells, [str(index), record, "", "", record_template or "按测试标准判定。", "□P □F", ""])


def add_compliance_table(document: Document, title: str, compliance_bug_info: str) -> None:
    table = document.add_table(rows=2, cols=5)
    table.style = "Table Grid"
    set_table_row(table.rows[0].cells, ["序号", "需求编号/DFMEA编号/风险管理编号", "需求描述", "测试结论", "备注"])
    set_table_row(table.rows[1].cells, ["1", "", title, compliance_bug_info or "待记录需求符合性结论。", ""])


def add_bug_table(document: Document) -> None:
    table = document.add_table(rows=2, cols=6)
    table.style = "Table Grid"
    set_table_row(table.rows[0].cells, ["序号", "问题描述", "涉及需求编号", "BUG编号（JIRA系统）", "RPN", "Bug解决状态"])
    set_table_row(table.rows[1].cells, ["1", "", "", "", "", ""])


def has_source_blocks(item) -> bool:
    return bool(item.source_blocks and item.source_type == "document") if hasattr(item, "source_type") else bool(item.source_blocks)


def append_source_blocks(document: Document, blocks: list[dict]) -> None:
    for block in blocks:
        if block.get("type") == "spacer":
            document.add_paragraph("")
            continue
        if block.get("type") == "table" and isinstance(block.get("rows"), list):
            add_source_table(document, block["rows"])
            continue
        text = str(block.get("text") or "").strip()
        if text:
            document.add_paragraph(text)


def add_source_table(document: Document, rows: list[list[str]]) -> None:
    usable_rows = [row for row in rows if row]
    if not usable_rows:
        return
    column_count = max(len(row) for row in usable_rows)
    table = document.add_table(rows=len(usable_rows), cols=column_count)
    table.style = "Table Grid"
    for row_index, row_values in enumerate(usable_rows):
        padded_values = [str(value) for value in row_values] + [""] * (column_count - len(row_values))
        set_table_row(table.rows[row_index].cells, padded_values)


def set_table_row(cells, values: list[str]) -> None:
    for cell, value in zip(cells, values):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.text = value
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            paragraph.paragraph_format.left_indent = Pt(0)
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.right_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)


def split_record_lines(record_template: str) -> list[str]:
    normalized = record_template.replace("；", "\n").replace(";", "\n")
    return [line.strip(" \t\r\n。.") for line in normalized.splitlines() if line.strip(" \t\r\n。.")]


def numbered_lines(values: list[str], fallback: str) -> str:
    usable_values = [value for value in values if value]
    if not usable_values:
        return fallback
    return "\n".join(f"{index}. {value}" for index, value in enumerate(usable_values, start=1))
